import pdfplumber
from docx import Document


def convert_pdf_to_word(pdf_path, output_docx):
    doc = Document()

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:

            # ---- TEXT (PARAGRAPHS) ----
            text = page.extract_text()
            if text:
                for line in text.split("\n"):
                    if line.strip():
                        p = doc.add_paragraph(line)
                        p.paragraph_format.space_after = 8

            # ---- TABLES ----
            tables = page.extract_tables()
            for table in tables:
                if not table or len(table) < 2:
                    continue

                rows = len(table)
                cols = len(table[0])

                word_table = doc.add_table(rows=rows, cols=cols)
                word_table.style = "Table Grid"

                for r in range(rows):
                    for c in range(cols):
                        cell_text = table[r][c] if table[r][c] else ""
                        word_table.rows[r].cells[c].text = str(cell_text)

            doc.add_page_break()

    doc.save(output_docx)
